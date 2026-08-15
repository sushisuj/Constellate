"""Tests for extract.py.

Covers a baseline case for every supported file type (so a regression
anywhere in extract_text() gets caught, not just in the paths touched
below) plus the two behaviours added on top of the ported Chatbot
pipeline: per-page table extraction and per-page OCR fallback.

PDF fixtures are built with reportlab at test time rather than committed
as binary files -- no binary fixtures in git, and what each test PDF
actually contains is readable directly from the test.
"""

import io

import pytest

from app import extract


def _pdf_with_text(lines):
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    for line in lines:
        c.drawString(100, 700, line)
        c.showPage()
    c.save()
    return buf.getvalue()


def _pdf_with_table():
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter)
    data = [["Quarter", "Revenue", "Profit"], ["Q1", "100", "20"], ["Q2", "150", "35"]]
    table = Table(data)
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 1, colors.black)]))
    doc.build([table])
    return buf.getvalue()


def _pdf_with_image_only_page(caption_text):
    """3-page PDF: real digital text, then a page that's only an image
    (the caption is drawn as pixels, not as text -- there's nothing for
    pypdf to extract on that page), then real digital text again."""
    from PIL import Image, ImageDraw
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas

    img = Image.new("RGB", (600, 200), "white")
    ImageDraw.Draw(img).text((20, 80), caption_text, fill="black")
    img_buf = io.BytesIO()
    img.save(img_buf, format="PNG")
    img_buf.seek(0)

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.drawString(100, 700, "Page one has real digital text in it.")
    c.showPage()
    c.drawImage(ImageReader(img_buf), 100, 500, width=400, height=133)
    c.showPage()
    c.drawString(100, 700, "Page three also has real digital text.")
    c.showPage()
    c.save()
    return buf.getvalue()


class TestExtractTextBaseline:
    def test_markdown_decodes_directly(self):
        text = extract.extract_text("notes.md", b"## Heading\nBody text.")
        assert text == "## Heading\nBody text."

    def test_txt_decodes_directly(self):
        text = extract.extract_text("notes.txt", "plain ascii text".encode("utf-8"))
        assert text == "plain ascii text"

    def test_unsupported_extension_raises(self):
        with pytest.raises(extract.UnsupportedFileType):
            extract.extract_text("archive.zip", b"whatever")

    def test_docx_extracts_paragraphs(self):
        import docx

        buf = io.BytesIO()
        document = docx.Document()
        document.add_paragraph("First paragraph.")
        document.add_paragraph("Second paragraph.")
        document.save(buf)
        text = extract.extract_text("notes.docx", buf.getvalue())
        assert "First paragraph." in text
        assert "Second paragraph." in text

    def test_image_with_no_readable_text_raises(self):
        from PIL import Image

        img = Image.new("RGB", (50, 50), "white")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        with pytest.raises(extract.UnsupportedFileType):
            extract.extract_text("blank.png", buf.getvalue())


class TestTableExtraction:
    def test_table_rows_and_columns_preserved(self):
        text = extract.extract_text("report.pdf", _pdf_with_table())
        assert "[Table]" in text
        assert "Quarter | Revenue | Profit" in text
        assert "Q1 | 100 | 20" in text
        assert "Q2 | 150 | 35" in text

    def test_plain_text_page_gets_no_table_block(self):
        text = extract.extract_text("memo.pdf", _pdf_with_text(["Just a memo, no table here."]))
        assert "[Table]" not in text
        assert "Just a memo" in text


class TestPerPageOCRFallback:
    @pytest.fixture
    def mixed_pdf_bytes(self):
        return _pdf_with_image_only_page("SCANNED DIAGRAM CONTENT")

    def test_image_only_page_gets_ocr_content(self, mixed_pdf_bytes):
        text = extract.extract_text("mixed.pdf", mixed_pdf_bytes)
        assert "SCANNED DIAGRAM CONTENT" in text.upper()

    def test_digital_text_pages_survive_alongside_ocr_page(self, mixed_pdf_bytes):
        text = extract.extract_text("mixed.pdf", mixed_pdf_bytes)
        assert "Page one has real digital text in it." in text
        assert "Page three also has real digital text." in text

    def test_ocr_only_runs_on_the_page_that_needs_it(self, monkeypatch, mixed_pdf_bytes):
        import pytesseract

        calls = []
        original = pytesseract.image_to_string

        def counting(image):
            calls.append(True)
            return original(image)

        monkeypatch.setattr(pytesseract, "image_to_string", counting)
        extract.extract_text("mixed.pdf", mixed_pdf_bytes)
        # 3 pages, only the middle one is image-only -- OCR should run
        # exactly once, not once per page and not zero times.
        assert len(calls) == 1

    def test_fully_digital_pdf_never_invokes_ocr(self, monkeypatch):
        import fitz

        def fail_if_called(*args, **kwargs):
            raise AssertionError("fitz.open should not run when no page needs OCR")

        monkeypatch.setattr(fitz, "open", fail_if_called)
        text = extract.extract_text("plain.pdf", _pdf_with_text(["All digital, no OCR needed."]))
        assert "All digital, no OCR needed." in text
